"""Reading a DOCX as plain text and optional chapter list.

Same contract as EPUB/PDF readers: no database, no model, no request state.
Chapters come from Word Heading 1 / Title-style paragraphs when there are at
least two; otherwise the file travels the plain-text path like a .txt upload.
"""

from typing import List, Optional, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def is_docx_filename(filename: str) -> bool:
    return filename.lower().endswith('.docx')


def extract_docx_book(
    filepath: str,
) -> Tuple[str, Optional[List[str]], Optional[str], Optional[str]]:
    """Read a DOCX and return (text, chapters, title, author).

    `chapters` is a list of plain-text chapter bodies when Heading 1 (or the
    document Title style used as a section break) appears at least twice;
    otherwise None, so the rest of the pipeline treats the file like .txt.
    """
    document = Document(filepath)
    paragraphs = []
    for paragraph in document.paragraphs:
        text = ' '.join(paragraph.text.split()).strip()
        if text:
            paragraphs.append((text, _is_chapter_heading(paragraph)))

    if not paragraphs:
        # Tables sometimes hold the only readable content (scripts, play text).
        for table in document.tables:
            for row in table.rows:
                cells = [' '.join(cell.text.split()).strip() for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    paragraphs.append((' | '.join(cells), False))

    title = None
    author = None
    props = document.core_properties
    if props is not None:
        title = (props.title or '').strip() or None
        author = (props.author or '').strip() or None

    heading_indexes = [i for i, (_, is_heading) in enumerate(paragraphs) if is_heading]
    chapters: Optional[List[str]] = None
    if len(heading_indexes) >= 2:
        chapters = []
        for index, start in enumerate(heading_indexes):
            end = heading_indexes[index + 1] if index + 1 < len(heading_indexes) else len(paragraphs)
            body = [paragraphs[i][0] for i in range(start, end)]
            chapters.append('\n\n'.join(body))

    text = '\n\n'.join(text for text, _ in paragraphs)
    return text, chapters, title, author


def _is_chapter_heading(paragraph) -> bool:
    """True for a Word paragraph that should open a new chapter."""
    style = paragraph.style
    if style is None:
        return False
    name = (style.name or '').strip().lower()
    if name in {'title', 'heading 1', 'heading1'}:
        return True
    # Localised Word builds rename styles; the outline level still marks H1.
    try:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.paragraph_format.outline_level == 0:
            # outline_level 0 is Heading 1; Title may also report it.
            if name.startswith('heading') or name == 'title':
                return True
    except (AttributeError, ValueError):
        pass
    return False
